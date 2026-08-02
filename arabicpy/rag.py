"""Local RAG index for built-in knowledge and user-uploaded documents."""

from __future__ import annotations

import math
import os
import re
import shutil
import hashlib
from collections import Counter
from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class KnowledgeChunk:
    title: str
    text: str


KNOWLEDGE = (
    KnowledgeChunk("الطباعة والمتغيرات", """لغة الباء تحول الكود العربي إلى Python. الإسناد يكتب هكذا: س = 5
الطباعة تستخدم اطبع مع أقواس: اطبع("مرحبا") أو اطبع(س). تدعم الأعداد والنصوص وعمليات + - * /."""),
    KnowledgeChunk("الشروط والمنطق", """الشرط يكتب:
اذا س > 5:
    اطبع("كبير")
الكلمات المنطقية المدعومة: و، او أو أو، ليس. القيم المنطقية: صح وخطأ. الفرع الآخر يستخدم والا."""),
    KnowledgeChunk("الحلقات", """حلقة لكل تمر على قائمة:
لكل عنصر في [1، 2، 3]:
    اطبع(عنصر)
توجد حلقة بينما. داخل الحلقات: توقف تعني break، استمر تعني continue، وتجاوز تعني pass."""),
    KnowledgeChunk("الدوال", """تعريف الدالة واستدعاؤها:
دالة جمع(أ، ب):
    ارجع أ + ب

اطبع(جمع(2، 3))
تستخدم ارجع لإرجاع قيمة، ويمكن استعمال الفاصلة العربية أو الإنجليزية بين المعاملات."""),
    KnowledgeChunk("القوائم والقواميس", """القائمة: ارقام = [10، 20، 30]
الوصول بالفهارس يبدأ من صفر: اطبع(ارقام[0]). يدعم الوصول المتسلسل مثل بيانات[0][1].
القاموس: بيانات = {"اسم": "سعود"، "عمر": 20}
والوصول: اطبع(بيانات["اسم"])."""),
    KnowledgeChunk("الإدخال والتحويل", """ادخل() يقرأ نصا من المستخدم. رقم(القيمة) يحول القيمة إلى رقم، ونوع(القيمة) يعيد نوعها."""),
    KnowledgeChunk("قواعد الدقة", """لا تخترع أمرا أو صيغة غير موثقة في لغة الباء. إذا لم تكن الميزة مدعومة فقل ذلك بوضوح. استخدم النقطتين بعد اذا ولكل وبينما ودالة، وأربع مسافات داخل الكتلة."""),
    KnowledgeChunk("تطبيقات Android", """بيئة الباء تستطيع إنشاء مشروع Android ومعاينة Kivy وتصدير main.py وbuildozer.spec. بناء APK يستخدم WSL2 وBuildozer، بينما تصدير المشروع لا يحتاج وجود WSL."""),
)


_ARABIC_MARKS = re.compile(r"[\u064b-\u065f\u0670]")
# Keep Arabic letters while excluding Arabic punctuation such as ؟ and ،.
_TOKEN = re.compile(r"[A-Za-z0-9_]+|[\u0621-\u063a\u0641-\u064a]+", re.UNICODE)


def _tokens(text: str) -> list[str]:
    text = _ARABIC_MARKS.sub("", str(text).lower())
    text = text.translate(str.maketrans({"أ": "ا", "إ": "ا", "آ": "ا", "ى": "ي", "ة": "ه"}))
    return _TOKEN.findall(text)


SUPPORTED_DOCUMENTS = {".txt", ".md", ".apy", ".py", ".json", ".csv", ".pdf", ".docx"}
_OCR_ENGINE = None


def documents_directory() -> Path:
    root = os.environ.get("LOCALAPPDATA")
    directory = (Path(root) if root else Path.home()) / "AlBaa" / "rag_documents"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def import_document(source) -> Path:
    """Copy a document into the persistent RAG library without trusting its name."""
    source = Path(source)
    extension = source.suffix.lower()
    if extension not in SUPPORTED_DOCUMENTS:
        raise ValueError("نوع المستند غير مدعوم")
    if not source.is_file() or source.stat().st_size > 25 * 1024 * 1024:
        raise ValueError("المستند غير موجود أو حجمه أكبر من 25MB")
    digest = hashlib.sha256(source.read_bytes()).hexdigest()[:12]
    safe_stem = re.sub(r"[^\w\u0600-\u06ff.-]+", "_", source.stem).strip("._") or "document"
    destination = documents_directory() / f"{safe_stem}-{digest}{extension}"
    if not destination.exists():
        shutil.copy2(source, destination)
    extracted = destination.with_suffix(destination.suffix + ".rag")
    if not extracted.exists():
        text = _read_document(destination)
        if not text.strip():
            raise ValueError("لم يمكن استخراج نص من المستند حتى باستخدام OCR")
        extracted.write_text(text, encoding="utf-8")
    return destination


def _read_document(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        scanned_pages = [index for index, text in enumerate(pages) if len(text.strip()) < 20]
        if scanned_pages:
            recognized = _ocr_pdf_pages(path, scanned_pages)
            for index, text in recognized.items():
                if text.strip():
                    pages[index] = text
        return "\n\n".join(pages)
    if extension == ".docx":
        from docx import Document
        return "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)
    return path.read_text(encoding="utf-8", errors="replace")


def _get_ocr_engine():
    """Load the Arabic/English mobile OCR model only when scanned pages exist."""
    global _OCR_ENGINE
    if _OCR_ENGINE is None:
        from paddleocr import PaddleOCR
        _OCR_ENGINE = PaddleOCR(
            lang="ar",
            ocr_version="PP-OCRv5",
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            enable_mkldnn=False,
        )
    return _OCR_ENGINE


def _ocr_result_text(results) -> str:
    lines = []
    for result in results or []:
        data = getattr(result, "json", result)
        if callable(data):
            data = data()
        if isinstance(data, dict):
            data = data.get("res", data)
            texts = data.get("rec_texts", [])
            lines.extend(str(text).strip() for text in texts if str(text).strip())
        elif isinstance(data, (list, tuple)):
            for item in data:
                if isinstance(item, (list, tuple)) and len(item) > 1:
                    value = item[1][0] if isinstance(item[1], (list, tuple)) else item[1]
                    if str(value).strip():
                        lines.append(str(value).strip())
    return "\n".join(lines)


def _ocr_pdf_pages(path: Path, page_indexes) -> dict[int, str]:
    """Render only image-only pages and recognize Arabic/English locally."""
    import pypdfium2 as pdfium
    import numpy as np
    engine = _get_ocr_engine()
    document = pdfium.PdfDocument(str(path))
    recognized = {}
    try:
        for index in page_indexes:
            page = document[index]
            try:
                image = page.render(scale=2.0).to_pil()
                recognized[index] = _ocr_result_text(engine.predict(np.asarray(image)))
            finally:
                page.close()
    finally:
        document.close()
    return recognized


def _split_document(title: str, text: str, size: int = 1200, overlap: int = 150):
    text = re.sub(r"\r\n?", "\n", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        end = min(len(text), start + size)
        if end < len(text):
            boundary = max(text.rfind("\n", start, end), text.rfind(". ", start, end))
            if boundary > start + size // 2:
                end = boundary + 1
        part = text[start:end].strip()
        if part:
            chunks.append(KnowledgeChunk(title, part))
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def user_knowledge() -> tuple[KnowledgeChunk, ...]:
    chunks = []
    for path in sorted(documents_directory().iterdir()):
        if path.is_file() and path.suffix.lower() in SUPPORTED_DOCUMENTS:
            try:
                extracted = path.with_suffix(path.suffix + ".rag")
                text = extracted.read_text(encoding="utf-8") if extracted.exists() else _read_document(path)
                chunks.extend(_split_document(path.stem.rsplit("-", 1)[0], text))
            except (OSError, ValueError, ImportError, RuntimeError):
                continue
    return tuple(chunks)


def retrieve(query: str, limit: int = 4) -> list[KnowledgeChunk]:
    """Return the most relevant verified chunks using a compact BM25-style score."""
    query_tokens = _tokens(query)
    if not query_tokens or limit <= 0:
        return []
    query_counts = Counter(query_tokens)
    documents = KNOWLEDGE + user_knowledge()
    document_tokens_list = tuple(_tokens(f"{chunk.title} {chunk.text}") for chunk in documents)
    document_frequency = Counter(token for tokens in map(set, document_tokens_list) for token in tokens)
    scores = []
    total = len(documents)
    for index, document_tokens in enumerate(document_tokens_list):
        counts = Counter(document_tokens)
        score = 0.0
        for token, query_weight in query_counts.items():
            frequency = counts[token]
            if not frequency:
                continue
            inverse_frequency = math.log(1 + (total + 1) / (1 + document_frequency[token]))
            score += query_weight * inverse_frequency * (1 + math.log(frequency))
        if score:
            scores.append((score, index))
    scores.sort(key=lambda item: (-item[0], item[1]))
    return [documents[index] for _score, index in scores[:limit]]


def context_for(query: str, limit: int = 4) -> str:
    chunks = retrieve(query, limit)
    if not chunks:
        return "لا توجد معرفة موثقة مرتبطة مباشرة بالسؤال. لا تخترع صيغة غير موجودة."
    return "\n\n".join(f"[{chunk.title}]\n{chunk.text}" for chunk in chunks)
